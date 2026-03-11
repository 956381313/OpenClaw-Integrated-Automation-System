#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档编辑脚本
支持编辑现有Word文档的各种操作
"""

import argparse
import os
import sys
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
import json

def edit_document_content(input_path, output_path, edits):
    """
    编辑文档内容
    
    Args:
        input_path: 输入文件路径
        output_path: 输出文件路径
        edits: 编辑操作列表
    """
    try:
        if not os.path.exists(input_path):
            print(f"输入文件不存在: {input_path}")
            return False
        
        doc = Document(input_path)
        
        for edit in edits:
            action = edit.get('action')
            
            if action == 'replace_text':
                # 替换文本
                old_text = edit.get('old_text')
                new_text = edit.get('new_text')
                if old_text and new_text is not None:
                    replace_text_in_document(doc, old_text, new_text)
            
            elif action == 'add_paragraph':
                # 添加段落
                text = edit.get('text', '')
                position = edit.get('position', 'end')  # 'start', 'end', or index
                add_paragraph(doc, text, position)
            
            elif action == 'remove_paragraph':
                # 删除段落
                pattern = edit.get('pattern')
                if pattern:
                    remove_paragraph_by_pattern(doc, pattern)
            
            elif action == 'update_heading':
                # 更新标题
                level = edit.get('level', 1)
                old_text = edit.get('old_text')
                new_text = edit.get('new_text')
                if old_text and new_text:
                    update_heading(doc, level, old_text, new_text)
            
            elif action == 'add_table':
                # 添加表格
                data = edit.get('data', [])
                if data:
                    add_table(doc, data, edit.get('position', 'end'))
            
            elif action == 'update_table':
                # 更新表格
                table_index = edit.get('table_index', 0)
                updates = edit.get('updates', {})
                update_table_cells(doc, table_index, updates)
            
            elif action == 'add_image':
                # 添加图片
                image_path = edit.get('image_path')
                width = edit.get('width', 4)  # 默认4英寸
                height = edit.get('height')
                if image_path and os.path.exists(image_path):
                    add_image(doc, image_path, width, height, edit.get('position', 'end'))
            
            elif action == 'apply_style':
                # 应用样式
                style_name = edit.get('style_name')
                pattern = edit.get('pattern')
                if style_name and pattern:
                    apply_style_to_text(doc, style_name, pattern)
        
        # 保存文档
        doc.save(output_path)
        print(f"文档已编辑并保存: {output_path}")
        return True
        
    except Exception as e:
        print(f"编辑文档时出错: {e}")
        return False

def replace_text_in_document(doc, old_text, new_text):
    """在文档中替换文本"""
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)
    
    # 替换表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_text in cell.text:
                    cell.text = cell.text.replace(old_text, new_text)

def add_paragraph(doc, text, position='end'):
    """添加段落"""
    if position == 'start':
        # 在文档开头添加
        new_paragraph = doc.add_paragraph(text)
        doc._body.insert(0, new_paragraph._element)
    elif position == 'end':
        # 在文档末尾添加
        doc.add_paragraph(text)
    elif isinstance(position, int):
        # 在指定位置添加
        if position < len(doc.paragraphs):
            new_paragraph = doc.add_paragraph(text)
            target_paragraph = doc.paragraphs[position]
            target_paragraph._element.addnext(new_paragraph._element)

def remove_paragraph_by_pattern(doc, pattern):
    """根据模式删除段落"""
    paragraphs_to_remove = []
    
    for i, paragraph in enumerate(doc.paragraphs):
        if re.search(pattern, paragraph.text):
            paragraphs_to_remove.append(i)
    
    # 从后往前删除，避免索引变化
    for i in sorted(paragraphs_to_remove, reverse=True):
        if i < len(doc.paragraphs):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)

def update_heading(doc, level, old_text, new_text):
    """更新标题"""
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith(f'Heading {level}') and old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)

def add_table(doc, data, position='end'):
    """添加表格"""
    if not data:
        return
    
    rows = len(data)
    cols = len(data[0]) if rows > 0 else 0
    
    if rows == 0 or cols == 0:
        return
    
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid'
    
    # 填充数据
    for i in range(rows):
        for j in range(cols):
            if i < len(data) and j < len(data[i]):
                table.cell(i, j).text = str(data[i][j])
    
    # 处理位置
    if position == 'start':
        # 移动到文档开头
        doc._body.insert(0, table._element)

def update_table_cells(doc, table_index, updates):
    """更新表格单元格"""
    if table_index >= len(doc.tables):
        print(f"表格索引 {table_index} 超出范围")
        return
    
    table = doc.tables[table_index]
    
    for cell_ref, new_value in updates.items():
        # 解析单元格引用，如 "A1", "B2"
        if len(cell_ref) >= 2:
            col_char = cell_ref[0].upper()
            row_str = cell_ref[1:]
            
            try:
                col_index = ord(col_char) - ord('A')
                row_index = int(row_str) - 1  # 转换为0-based索引
                
                if 0 <= row_index < len(table.rows) and 0 <= col_index < len(table.columns):
                    table.cell(row_index, col_index).text = str(new_value)
            except (ValueError, IndexError):
                print(f"无效的单元格引用: {cell_ref}")

def add_image(doc, image_path, width, height=None, position='end'):
    """添加图片"""
    from docx.shared import Inches
    
    if height:
        doc.add_picture(image_path, width=Inches(width), height=Inches(height))
    else:
        doc.add_picture(image_path, width=Inches(width))
    
    # 处理位置
    if position == 'start':
        # 获取最后添加的段落（包含图片）
        last_paragraph = doc.paragraphs[-1]
        doc._body.insert(0, last_paragraph._element)

def apply_style_to_text(doc, style_name, pattern):
    """应用样式到匹配的文本"""
    # 首先确保样式存在
    if style_name not in doc.styles:
        print(f"样式不存在: {style_name}")
        return
    
    style = doc.styles[style_name]
    
    for paragraph in doc.paragraphs:
        if re.search(pattern, paragraph.text):
            paragraph.style = style

def bulk_replace(input_path, output_path, replacements):
    """
    批量替换文本
    
    Args:
        input_path: 输入文件路径
        output_path: 输出文件路径
        replacements: 替换字典 {旧文本: 新文本}
    """
    try:
        if not os.path.exists(input_path):
            print(f"输入文件不存在: {input_path}")
            return False
        
        doc = Document(input_path)
        
        # 替换段落中的文本
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)
        
        # 替换表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for old_text, new_text in replacements.items():
                        if old_text in cell.text:
                            cell.text = cell.text.replace(old_text, new_text)
        
        doc.save(output_path)
        print(f"批量替换完成: {output_path}")
        return True
        
    except Exception as e:
        print(f"批量替换时出错: {e}")
        return False

def format_document(input_path, output_path, formatting_rules):
    """
    格式化文档
    
    Args:
        input_path: 输入文件路径
        output_path: 输出文件路径
        formatting_rules: 格式化规则
    """
    try:
        if not os.path.exists(input_path):
            print(f"输入文件不存在: {input_path}")
            return False
        
        doc = Document(input_path)
        
        # 应用格式化规则
        for rule in formatting_rules:
            rule_type = rule.get('type')
            
            if rule_type == 'font':
                # 字体格式化
                apply_font_formatting(doc, rule)
            
            elif rule_type == 'paragraph':
                # 段落格式化
                apply_paragraph_formatting(doc, rule)
            
            elif rule_type == 'table':
                # 表格格式化
                apply_table_formatting(doc, rule)
        
        doc.save(output_path)
        print(f"文档格式化完成: {output_path}")
        return True
        
    except Exception as e:
        print(f"格式化文档时出错: {e}")
        return False

def apply_font_formatting(doc, rule):
    """应用字体格式化"""
    pattern = rule.get('pattern', '')
    font_name = rule.get('font_name')
    font_size = rule.get('font_size')
    bold = rule.get('bold')
    italic = rule.get('italic')
    color = rule.get('color')
    
    for paragraph in doc.paragraphs:
        if not pattern or re.search(pattern, paragraph.text):
            for run in paragraph.runs:
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = Pt(font_size)
                if bold is not None:
                    run.bold = bold
                if italic is not None:
                    run.italic = italic
                if color:
                    if color.startswith('#'):
                        color = color.lstrip('#')
                        rgb = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
                        run.font.color.rgb = RGBColor(*rgb)

def apply_paragraph_formatting(doc, rule):
    """应用段落格式化"""
    pattern = rule.get('pattern', '')
    alignment = rule.get('alignment')
    line_spacing = rule.get('line_spacing')
    space_before = rule.get('space_before')
    space_after = rule.get('space_after')
    
    alignment_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    
    for paragraph in doc.paragraphs:
        if not pattern or re.search(pattern, paragraph.text):
            if alignment and alignment in alignment_map:
                paragraph.alignment = alignment_map[alignment]
            if line_spacing:
                paragraph.paragraph_format.line_spacing = line_spacing
            if space_before:
                paragraph.paragraph_format.space_before = Pt(space_before)
            if space_after:
                paragraph.paragraph_format.space_after = Pt(space_after)

def apply_table_formatting(doc, rule):
    """应用表格格式化"""
    table_index = rule.get('table_index', 0)
    if table_index >= len(doc.tables):
        return
    
    table = doc.tables[table_index]
    style = rule.get('style')
    
    if style:
        table.style = style

def main():
    parser = argparse.ArgumentParser(description='编辑Word文档')
    parser.add_argument('--input', '-i', required=True, help='输入文件路径')
    parser.add_argument('--output', '-o', required=True, help='输出文件路径')
    parser.add_argument('--action', '-a', default='edit',
                       choices=['edit', 'replace', 'format'],
                       help='编辑操作类型')
    parser.add_argument('--edits', help='编辑操作JSON（用于edit类型）')
    parser.add_argument('--replacements', help='替换字典JSON（用于replace类型）')
    parser.add_argument('--formatting', help='格式化规则JSON（用于format类型）')
    
    args = parser.parse_args()
    
    if not os.path.exists(args.input):
        print(f"输入文件不存在: {args.input}")
        sys.exit(1)
    
    if args.action == 'edit':
        if not args.edits:
            print("错误：edit类型需要--edits参数")
            sys.exit(1)
        
        try:
            edits = json.loads(args.edits)
        except json.JSONDecodeError as e:
            print(f"解析edits JSON时出错: {e}")
            sys.exit(1)
        
        success = edit_document_content(args.input, args.output, edits)
    
    elif args.action == 'replace':
        if not args.replacements:
            print("错误：replace类型需要--replacements参数")
            sys.exit(1)
        
        try:
            replacements = json.loads(args.replacements)
        except json.JSONDecodeError as e:
            print(f"解析replacements JSON时出错: {e}")
            sys.exit(1)
        
        success = bulk_replace(args.input, args.output, replacements)
    
    elif args.action == 'format':
        if not args.formatting:
            print("错误：format类型需要--formatting参数")
            sys.exit(1)
        
        try:
            formatting_rules = json.loads(args.formatting)
        except json.JSONDecodeError as e:
            print(f"解析formatting JSON时出错: {e}")
            sys.exit(1)
        
        success = format_document(args.input, args.output, formatting_rules)
    
    else:
        print(f"不支持的编辑操作: {args.action}")
        sys.exit(1)
    
    if not success:
        sys.exit(1)

if __name__ == "__main__":
    main()