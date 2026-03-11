#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档创建脚本
支持创建各种类型的Word文档，包括中文文档
"""

import argparse
import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import json

def create_basic_document(output_path, title="新文档", content=None):
    """
    创建基本文档
    
    Args:
        output_path: 输出文件路径
        title: 文档标题
        content: 文档内容列表，每个元素是一个段落
    """
    try:
        # 创建文档对象
        doc = Document()
        
        # 添加标题
        doc.add_heading(title, 0)
        
        # 添加创建时间
        current_time = datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")
        doc.add_paragraph(f"创建时间: {current_time}")
        
        # 添加内容
        if content:
            for item in content:
                if isinstance(item, dict):
                    # 带格式的段落
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(item.get('text', ''))
                    
                    # 设置字体格式
                    if 'bold' in item and item['bold']:
                        run.bold = True
                    if 'italic' in item and item['italic']:
                        run.italic = True
                    if 'underline' in item and item['underline']:
                        run.underline = True
                    if 'font_size' in item:
                        run.font.size = Pt(item['font_size'])
                    if 'font_color' in item:
                        color = item['font_color']
                        if isinstance(color, str) and color.startswith('#'):
                            # 十六进制颜色
                            color = color.lstrip('#')
                            rgb = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
                            run.font.color.rgb = RGBColor(*rgb)
                else:
                    # 普通段落
                    doc.add_paragraph(str(item))
        
        # 保存文档
        doc.save(output_path)
        print(f"文档已创建: {output_path}")
        return True
        
    except Exception as e:
        print(f"创建文档时出错: {e}")
        return False

def create_chinese_document(output_path, title="中文文档", content=None):
    """
    创建中文文档，使用中文字体
    
    Args:
        output_path: 输出文件路径
        title: 文档标题
        content: 文档内容
    """
    try:
        doc = Document()
        
        # 设置中文字体（如果系统中有这些字体）
        # 注意：这需要系统中已安装相应字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'  # 或 '宋体', '黑体'
        font.size = Pt(12)
        
        # 添加中文标题
        doc.add_heading(title, 0)
        
        # 添加中文内容
        if content:
            for paragraph_text in content:
                p = doc.add_paragraph(paragraph_text)
                # 设置段落格式
                p.paragraph_format.line_spacing = 1.5  # 1.5倍行距
                p.paragraph_format.space_after = Pt(12)  # 段后间距
        
        # 添加中文页脚
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"第 1 页"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.save(output_path)
        print(f"中文文档已创建: {output_path}")
        return True
        
    except Exception as e:
        print(f"创建中文文档时出错: {e}")
        return False

def create_from_template(template_path, output_path, data_dict):
    """
    基于模板创建文档
    
    Args:
        template_path: 模板文件路径
        output_path: 输出文件路径
        data_dict: 替换数据的字典
    """
    try:
        if not os.path.exists(template_path):
            print(f"模板文件不存在: {template_path}")
            return False
            
        doc = Document(template_path)
        
        # 替换文档中的占位符
        for paragraph in doc.paragraphs:
            for key, value in data_dict.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))
        
        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data_dict.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))
        
        doc.save(output_path)
        print(f"基于模板的文档已创建: {output_path}")
        return True
        
    except Exception as e:
        print(f"基于模板创建文档时出错: {e}")
        return False

def create_report_document(output_path, report_data):
    """
    创建报告文档
    
    Args:
        output_path: 输出文件路径
        report_data: 报告数据字典
    """
    try:
        doc = Document()
        
        # 报告标题
        doc.add_heading(report_data.get('title', '报告'), 0)
        
        # 报告信息
        info = report_data.get('info', {})
        if info:
            info_table = doc.add_table(rows=len(info)+1, cols=2)
            info_table.style = 'Light Grid'
            
            # 表头
            info_table.cell(0, 0).text = '项目'
            info_table.cell(0, 1).text = '内容'
            
            # 数据行
            for i, (key, value) in enumerate(info.items(), 1):
                info_table.cell(i, 0).text = str(key)
                info_table.cell(i, 1).text = str(value)
        
        # 报告正文
        content = report_data.get('content', [])
        for section in content:
            if 'heading' in section:
                doc.add_heading(section['heading'], level=section.get('level', 1))
            if 'text' in section:
                doc.add_paragraph(section['text'])
            if 'table' in section:
                table_data = section['table']
                if table_data:
                    rows = len(table_data)
                    cols = len(table_data[0]) if rows > 0 else 0
                    if rows > 0 and cols > 0:
                        table = doc.add_table(rows=rows, cols=cols)
                        table.style = 'Light Grid'
                        for i, row in enumerate(table_data):
                            for j, cell in enumerate(row):
                                table.cell(i, j).text = str(cell)
        
        # 结论
        conclusion = report_data.get('conclusion', '')
        if conclusion:
            doc.add_heading('结论', 2)
            doc.add_paragraph(conclusion)
        
        doc.save(output_path)
        print(f"报告文档已创建: {output_path}")
        return True
        
    except Exception as e:
        print(f"创建报告文档时出错: {e}")
        return False

def main():
    parser = argparse.ArgumentParser(description='创建Word文档')
    parser.add_argument('--output', '-o', required=True, help='输出文件路径')
    parser.add_argument('--type', '-t', default='basic', 
                       choices=['basic', 'chinese', 'template', 'report'],
                       help='文档类型')
    parser.add_argument('--title', help='文档标题')
    parser.add_argument('--content', help='文档内容（JSON格式）')
    parser.add_argument('--template', help='模板文件路径（用于template类型）')
    parser.add_argument('--data', help='数据文件路径（JSON格式）')
    
    args = parser.parse_args()
    
    # 处理内容
    content_data = None
    if args.content:
        try:
            content_data = json.loads(args.content)
        except json.JSONDecodeError:
            # 如果不是JSON，作为普通文本处理
            content_data = [args.content]
    
    # 根据类型创建文档
    if args.type == 'basic':
        success = create_basic_document(
            args.output,
            title=args.title or "新文档",
            content=content_data
        )
    
    elif args.type == 'chinese':
        success = create_chinese_document(
            args.output,
            title=args.title or "中文文档",
            content=content_data
        )
    
    elif args.type == 'template':
        if not args.template:
            print("错误：模板类型需要--template参数")
            sys.exit(1)
        
        data_dict = {}
        if args.data:
            try:
                with open(args.data, 'r', encoding='utf-8') as f:
                    data_dict = json.load(f)
            except Exception as e:
                print(f"读取数据文件时出错: {e}")
                sys.exit(1)
        
        success = create_from_template(
            args.template,
            args.output,
            data_dict
        )
    
    elif args.type == 'report':
        report_data = {}
        if args.data:
            try:
                with open(args.data, 'r', encoding='utf-8') as f:
                    report_data = json.load(f)
            except Exception as e:
                print(f"读取报告数据时出错: {e}")
                sys.exit(1)
        
        success = create_report_document(args.output, report_data)
    
    else:
        print(f"不支持的文档类型: {args.type}")
        sys.exit(1)
    
    if not success:
        sys.exit(1)

if __name__ == "__main__":
    main()