#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档格式转换脚本
支持Word文档与其他格式之间的转换
"""

import argparse
import os
import sys
import subprocess
from pathlib import Path
import tempfile

def word_to_pdf(input_path, output_path):
    """
    Word转PDF
    
    Args:
        input_path: Word文档路径
        output_path: PDF输出路径
    """
    try:
        # 方法1: 使用python-docx和reportlab（简单但功能有限）
        try:
            from docx import Document
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.units import inch
            
            doc = Document(input_path)
            c = canvas.Canvas(output_path, pagesize=letter)
            
            y = 750  # 起始Y坐标
            for paragraph in doc.paragraphs:
                text = paragraph.text
                if text.strip():  # 只处理非空段落
                    c.drawString(50, y, text[:100])  # 简单示例，实际需要更复杂的布局
                    y -= 20
                    if y < 50:  # 换页
                        c.showPage()
                        y = 750
            
            c.save()
            print(f"Word转PDF完成（简单方法）: {output_path}")
            return True
            
        except ImportError:
            print("reportlab未安装，尝试其他方法")
        
        # 方法2: 使用LibreOffice（推荐，需要安装LibreOffice）
        if sys.platform.startswith('win'):
            # Windows
            libreoffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
            if os.path.exists(libreoffice_path):
                cmd = [
                    libreoffice_path,
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', os.path.dirname(output_path),
                    input_path
                ]
                result = subprocess.run(cmd, capture_output=True, text=True)
                if result.returncode == 0:
                    # LibreOffice输出PDF到相同目录，重命名为目标路径
                    base_name = os.path.splitext(os.path.basename(input_path))[0]
                    temp_pdf = os.path.join(os.path.dirname(output_path), f"{base_name}.pdf")
                    if os.path.exists(temp_pdf):
                        os.rename(temp_pdf, output_path)
                        print(f"Word转PDF完成（LibreOffice）: {output_path}")
                        return True
            else:
                print("LibreOffice未找到，请安装LibreOffice或使用其他方法")
        
        # 方法3: 使用python-docx2pdf（需要安装）
        try:
            from docx2pdf import convert
            convert(input_path, output_path)
            print(f"Word转PDF完成（docx2pdf）: {output_path}")
            return True
        except ImportError:
            print("docx2pdf未安装，请安装: pip install docx2pdf")
        
        print("所有转换方法都失败，请检查依赖")
        return False
        
    except Exception as e:
        print(f"Word转PDF时出错: {e}")
        return False

def pdf_to_word(input_path, output_path):
    """
    PDF转Word
    
    Args:
        input_path: PDF文件路径
        output_path: Word文档输出路径
    """
    try:
        # 方法1: 使用pdf2docx（推荐）
        try:
            from pdf2docx import Converter
            
            cv = Converter(input_path)
            cv.convert(output_path, start=0, end=None)
            cv.close()
            print(f"PDF转Word完成（pdf2docx）: {output_path}")
            return True
            
        except ImportError:
            print("pdf2docx未安装，请安装: pip install pdf2docx")
        
        # 方法2: 使用PyPDF2和python-docx（基本文本提取）
        try:
            import PyPDF2
            from docx import Document
            
            doc = Document()
            
            with open(input_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    
                    if text.strip():
                        # 添加分页符（第一页除外）
                        if page_num > 0:
                            doc.add_page_break()
                        
                        # 添加文本
                        doc.add_paragraph(text)
            
            doc.save(output_path)
            print(f"PDF转Word完成（文本提取）: {output_path}")
            return True
            
        except ImportError:
            print("PyPDF2未安装，请安装: pip install PyPDF2")
        
        print("所有转换方法都失败，请检查依赖")
        return False
        
    except Exception as e:
        print(f"PDF转Word时出错: {e}")
        return False

def word_to_html(input_path, output_path):
    """
    Word转HTML
    
    Args:
        input_path: Word文档路径
        output_path: HTML输出路径
    """
    try:
        from docx import Document
        import html
        
        doc = Document(input_path)
        
        html_content = []
        html_content.append('<!DOCTYPE html>')
        html_content.append('<html lang="zh-CN">')
        html_content.append('<head>')
        html_content.append('<meta charset="UTF-8">')
        html_content.append('<meta name="viewport" content="width=device-width, initial-scale=1.0">')
        html_content.append(f'<title>{os.path.basename(input_path)}</title>')
        html_content.append('<style>')
        html_content.append('body { font-family: Arial, sans-serif; line-height: 1.6; margin: 20px; }')
        html_content.append('h1 { color: #333; }')
        html_content.append('p { margin-bottom: 10px; }')
        html_content.append('table { border-collapse: collapse; width: 100%; }')
        html_content.append('th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }')
        html_content.append('th { background-color: #f2f2f2; }')
        html_content.append('img { max-width: 100%; height: auto; }')
        html_content.append('</style>')
        html_content.append('</head>')
        html_content.append('<body>')
        
        for element in doc.element.body:
            if element.tag.endswith('p'):  # 段落
                paragraph = doc.paragraphs[0]  # 简化处理
                text = paragraph.text
                if text.strip():
                    # 检查是否是标题
                    if paragraph.style.name.startswith('Heading'):
                        level = int(paragraph.style.name.split()[-1])
                        html_content.append(f'<h{level}>{html.escape(text)}</h{level}>')
                    else:
                        html_content.append(f'<p>{html.escape(text)}</p>')
            
            elif element.tag.endswith('tbl'):  # 表格
                html_content.append('<table>')
                # 简化处理，实际需要更复杂的表格解析
                html_content.append('<tr><td>表格内容</td></tr>')
                html_content.append('</table>')
        
        html_content.append('</body>')
        html_content.append('</html>')
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(html_content))
        
        print(f"Word转HTML完成: {output_path}")
        return True
        
    except Exception as e:
        print(f"Word转HTML时出错: {e}")
        return False

def word_to_markdown(input_path, output_path):
    """
    Word转Markdown
    
    Args:
        input_path: Word文档路径
        output_path: Markdown输出路径
    """
    try:
        from docx import Document
        
        doc = Document(input_path)
        
        md_lines = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # 处理标题
            if paragraph.style.name.startswith('Heading'):
                level = int(paragraph.style.name.split()[-1])
                md_lines.append(f"{'#' * level} {text}")
            else:
                md_lines.append(text)
            
            md_lines.append('')  # 空行分隔
        
        # 处理表格（简化）
        for table in doc.tables:
            md_lines.append('')  # 空行
            # 这里可以添加更复杂的表格转换逻辑
            md_lines.append('| 表格 | 内容 |')
            md_lines.append('|------|------|')
            for row in table.rows:
                row_cells = [cell.text for cell in row.cells]
                md_lines.append(f"| {' | '.join(row_cells)} |")
            md_lines.append('')
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(md_lines))
        
        print(f"Word转Markdown完成: {output_path}")
        return True
        
    except Exception as e:
        print(f"Word转Markdown时出错: {e}")
        return False

def word_to_text(input_path, output_path):
    """
    Word转纯文本
    
    Args:
        input_path: Word文档路径
        output_path: 文本输出路径
    """
    try:
        # 方法1: 使用python-docx
        from docx import Document
        
        doc = Document(input_path)
        
        text_lines = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_lines.append(text)
        
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_lines.append('\t'.join(row_text))
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(text_lines))
        
        print(f"Word转文本完成: {output_path}")
        return True
        
    except Exception as e:
        print(f"Word转文本时出错: {e}")
        return False

def html_to_word(input_path, output_path):
    """
    HTML转Word
    
    Args:
        input_path: HTML文件路径
        output_path: Word文档输出路径
    """
    try:
        from docx import Document
        from bs4 import BeautifulSoup
        
        with open(input_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        doc = Document()
        
        # 处理标题
        for h_tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            level = int(h_tag.name[1])
            doc.add_heading(h_tag.get_text(), level)
        
        # 处理段落
        for p_tag in soup.find_all('p'):
            text = p_tag.get_text().strip()
            if text:
                doc.add_paragraph(text)
        
        # 处理列表
        for ul_tag in soup.find_all('ul'):
            for li_tag in ul_tag.find_all('li'):
                doc.add_paragraph(li_tag.get_text(), style='List Bullet')
        
        for ol_tag in soup.find_all('ol'):
            for li_tag in ol_tag.find_all('li'):
                doc.add_paragraph(li_tag.get_text(), style='List Number')
        
        # 处理表格（简化）
        for table_tag in soup.find_all('table'):
            # 这里可以添加更复杂的表格处理逻辑
            doc.add_paragraph("[表格内容]")
        
        doc.save(output_path)
        print(f"HTML转Word完成: {output_path}")
        return True
        
    except Exception as e:
        print(f"HTML转Word时出错: {e}")
        return False

def markdown_to_word(input_path, output_path):
    """
    Markdown转Word
    
    Args:
        input_path: Markdown文件路径
        output_path: Word文档输出路径
    """
    try:
        from docx import Document
        import re
        
        with open(input_path, 'r', encoding='utf-8') as f:
            md_lines = f.readlines()
        
        doc = Document()
        
        for line in md_lines:
            line = line.rstrip('\n')
            
            # 处理标题
            match = re.match(r'^(#+)\s+(.+)$', line)
            if match:
                level = len(match.group(1))
                text = match.group(2)
                doc.add_heading(text, level)
                continue
            
            # 处理列表
            if re.match(r'^[\-\*\+]\s+.+$', line):
                text = re.sub(r'^[\-\*\+]\s+', '', line)
                doc.add_paragraph(text, style='List Bullet')
                continue
            
            if re.match(r'^\d+\.\s+.+$', line):
                text = re.sub(r'^\d+\.\s+', '', line)
                doc.add_paragraph(text, style='List Number')
                continue
            
            # 处理表格（简化）
            if '|' in line and re.search(r'[a-zA-Z]', line):
                # 这里可以添加更复杂的表格处理逻辑
                continue
            
            # 普通段落
            if line.strip():
                doc.add_paragraph(line)
        
        doc.save(output_path)
        print(f"Markdown转Word完成: {output_path}")
        return True
        
    except Exception as e:
        print(f"Markdown转Word时出错: {e}")
        return False

def text_to_word(input_path, output_path):
    """
    纯文本转Word
    
    Args:
        input_path: 文本文件路径
        output_path: Word文档输出路径
    """
    try:
        from docx import Document
        
        with open(input_path, 'r', encoding='utf-8') as f:
            text_lines = f.readlines()
        
        doc = Document()
        
        for line in text_lines:
            line = line.rstrip('\n')
            if line.strip():
                doc.add_paragraph(line)
        
        doc.save(output_path)
        print(f"文本转Word完成: {output_path}")
        return True
        
    except Exception as e:
        print(f"文本转Word时出错: {e}")
        return False

def convert_format(input_path, output_path, format_type):
    """
    通用格式转换函数
    
    Args:
        input_path: 输入文件路径
        output_path: 输出文件路径
        format_type: 转换类型
    """
    # 检查输入文件
    if not os.path.exists(input_path):
        print(f"输入文件不存在: {input_path}")
        return False
    
    # 根据转换类型调用相应的函数
    if format_type == 'word_to_pdf':
        return word_to_pdf(input_path, output_path)
    
    elif format_type == 'pdf_to_word':
        return pdf_to_word(input_path, output_path)
    
    elif format_type == 'word_to_html':
        return word_to_html(input_path, output_path)
    
    elif format_type == 'html_to_word':
        return html_to_word(input_path, output_path)
    
    elif format_type == 'word_to_markdown':
        return word_to_markdown(input_path, output_path)
    
    elif format_type == 'markdown_to_word':
        return markdown_to_word(input_path, output_path)
    
    elif format_type == 'word_to_text':
        return word_to_text(input_path, output_path)
    
    elif format_type == 'text_to_word':
        return text_to_word(input_path, output_path)
    
    else:
        print(f"不支持的转换类型: {format_type}")
        return False

def main():
    parser = argparse.ArgumentParser(description='文档格式转换')
    parser.add_argument('--input', '-i', required=True, help='输入文件路径')
    parser.add_argument('--output', '-o', required=True, help='输出文件路径')
    parser.add_argument('--type', '-t', required=True,
                       choices=['word_to_pdf', 'pdf_to_word',
                                'word_to_html', 'html_to_word',
                                'word_to_markdown', 'markdown_to_word',
                                'word_to_text', 'text_to_word'],
                       help='转换类型')
    
    args = parser.parse_args()
    
    # 检查输入文件
    if not os.path.exists(args.input):
        print(f"输入文件不存在: {args.input}")
        sys.exit(1)
    
    # 执行转换
    success = convert_format(args.input, args.output, args.type)
    
    if not success:
        sys.exit(1)

if __name__ == "__main__":
    main()